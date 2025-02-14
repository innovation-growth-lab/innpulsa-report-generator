# Imports
import io
import asyncio
import streamlit as st
from docx import Document
from src.config.sections import sections_config
from src.data.process import aggregate_data
from src.data.loaders import load_data
from src.utils.output import generate_json_output
from src.services.feedback import save_suggestion
from src.services.openai_helpers import (
    generate_section_contents,
    generate_executive_summary,
    edit_report_sections,
)

# Initialize session state
if "report_generated" not in st.session_state:
    st.session_state.report_generated = False
    st.session_state.report_finalized = False
    st.session_state.json_str = None
    st.session_state.markdown_content = None
    st.session_state.resumen_ejecutivo = None
    st.session_state.edited_output = None
    st.session_state.report_sections = None
    st.session_state.success_message = None

# Page config
st.set_page_config(
    layout="centered", page_title="Generador de Reportes ZASCA", page_icon="📊"
)
st.image("assets/zasca_logo.png")

# Application Title
st.title("\U0001F4DA Generador de Reportes ZASCA")

tab1, tab2, tab3, tab4 = st.tabs(
    [
        "📝 Descripción General",
        "📋 Proceso de Uso",
        "⚠️ Consideraciones",
        "💡 Sugerencias",
    ]
)

with tab1:
    st.markdown(
        """
        <div style="background-color: #f9f9f9; padding: 20px; border-radius: 10px;">
            <div style="margin-bottom: 20px;">
                <p>
                    Esta herramienta está diseñada para ayudar a <strong>INNPULSA</strong> a 
                    crear los informes de cierre de los centros <strong>ZASCA</strong> que 
                    apoyan la estrategia de reindustrialización en Colombia.
                </p>
            </div>
        </div>
    """,
        unsafe_allow_html=True,
    )

    # add a space
    st.markdown(
        """
        <div style="height: 20px;"></div>
    """,
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns(2)

    with col1:
        st.markdown(
            """
            <div style="background-color: #f9f9f9; padding: 20px; border-radius: 10px; height: 100%;">
                <h5 style="color: #2c3e50;">✅ Lo que hace</h5>
                <ul style="margin-left: 20px;">
                    <li>Genera un primer borrador estructurado que sirve como base para el informe final</li>
                    <li>Calcula automáticamente indicadores clave a partir de los datos</li>
                    <li>Proporciona interpretaciones preliminares de los cambios observados</li>
                    <li>Estructura la información en secciones temáticas coherentes</li>
                    <li>Genera un resumen ejecutivo inicial</li>
                </ul>
            </div>
        """,
            unsafe_allow_html=True,
        )

    with col2:
        st.markdown(
            """
            <div style="background-color: #f9f9f9; padding: 20px; border-radius: 10px; height: 100%;">
                <h5 style="color: #2c3e50;">❌ Lo que NO hace</h5>
                <ul style="margin-left: 20px;">
                    <li>No produce un informe final listo para publicar</li>
                    <li>No reemplaza el análisis experto del equipo técnico</li>
                    <li>No genera el documento en formato PDF o Word</li>
                    <li>No valida la calidad o coherencia de los datos de entrada</li>
                    <li>No incorpora automáticamente información contextual no proporcionada</li>
                </ul>
            </div>
        """,
            unsafe_allow_html=True,
        )


with tab2:
    st.markdown(
        """
        <div style="background-color: #f9f9f9; padding: 20px; border-radius: 10px;">
            <ol>
                <li><strong>Preparación de Datos</strong>
                    <ul>
                        <li>Cargar el archivo Excel con los datos de diagnóstico y cierre</li>
                        <li>Revisar las variables disponibles en el menú lateral</li>
                        <li>Verificar qué variables están ausentes o incompletas</li>
                    </ul>
                </li>
                <li><strong>Generación del Reporte</strong>
                    <ul>
                        <li>Completar la información contextual del centro ZASCA</li>
                        <li>Seleccionar el modelo de IA según necesidad de precisión</li>
                        <li>Generar el reporte inicial</li>
                    </ul>
                </li>
                <li><strong>Revisión y Edición</strong>
                    <ul>
                        <li>Verificar la precisión de los cálculos e interpretaciones</li>
                        <li>Evaluar la relevancia de las variables incluidas</li>
                        <li>Considerar el contexto específico del centro</li>
                    </ul>
                </li>
            </ol>
        </div>
    """,
        unsafe_allow_html=True,
    )

with tab3:
    st.markdown(
        """
        <div style="background-color: #f9f9f9; padding: 20px; border-radius: 10px;">
            <div style="margin-bottom: 20px;">
                <h4 style="color: #2c3e50;">🔍 Verificación de Datos</h4>
                <p style="margin-left: 20px;">
                    El reporte generado debe ser revisado cuidadosamente. Los cálculos son 
                    automáticos pero requieren validación humana para asegurar su precisión
                    en el contexto específico.
                </p>
            </div> 
            <div style="margin-bottom: 20px;">
                <h4 style="color: #2c3e50;">💡 Interpretación Contextual</h4>
                <p style="margin-left: 20px;">
                    Las interpretaciones generadas son preliminares y deben ser ajustadas 
                    considerando:
                </p>
                <ul style="margin-left: 40px;">
                    <li>El contexto específico del centro</li>
                    <li>Las características del sector</li>
                    <li>La realidad socioeconómica local</li>
                </ul>
            </div>
            <div style="margin-bottom: 20px;">
                <h4 style="color: #2c3e50;">⚠️ Variables Faltantes</h4>
                <p style="margin-left: 20px;">
                    Cuando hay variables no encontradas (listadas en el menú lateral):
                </p>
                <ul style="margin-left: 40px;">
                    <li>Evaluar si son cruciales para el análisis</li>
                    <li>Considerar la necesidad de recolectar datos adicionales</li>
                    <li>Documentar las razones de su ausencia</li>
                </ul>
            </div>            
            <div style="margin-bottom: 20px;">
                <h4 style="color: #2c3e50;">✏️ Personalización Necesaria</h4>
                <p style="margin-left: 20px;">
                    El reporte debe ser personalizado para reflejar:
                </p>
                <ul style="margin-left: 40px;">
                    <li>Particularidades del sector y región</li>
                    <li>Objetivos específicos del centro</li>
                    <li>Desafíos únicos enfrentados</li>
                    <li>Contexto socioeconómico local</li>
                    <li>Factores externos relevantes</li>
                </ul>
            </div>
            <div style="margin-bottom: 20px;">
                <h4 style="color: #2c3e50;">📋 Resultado Final</h4>
                <p style="margin-left: 20px;">
                    El producto generado es un borrador estructurado que requiere:
                </p>
                <ul style="margin-left: 40px;">
                    <li>Revisión técnica del equipo</li>
                    <li>Validación de interpretaciones</li>
                    <li>Enriquecimiento con conocimiento local</li>
                    <li>Ajustes finales de estilo y formato</li>
                </ul>
            </div>
        </div>
    """,
        unsafe_allow_html=True,
    )

with tab4:
    st.markdown("### Sugerencias y Solicitudes")
    st.markdown(
        """
        <div style="background-color: #f9f9f9; padding: 20px; border-radius: 10px;">
            <p>¿Tienes alguna sugerencia para mejorar la herramienta o necesitas alguna funcionalidad adicional?
            Compártela con nosotros.</p>
        </div>
    """,
        unsafe_allow_html=True,
    )

    category = st.selectbox(
        "Categoría",
        ["Mejora de funcionalidad", "Nueva característica", "Reporte de error", "Otro"],
    )

    suggestion = st.text_area(
        "Sugerencia", placeholder="Describe tu sugerencia o solicitud..."
    )

    details = st.text_area(
        "Detalles adicionales",
        placeholder="Proporciona cualquier detalle adicional que nos ayude a entender mejor tu solicitud...",
    )

    if st.button("Enviar sugerencia"):
        if suggestion.strip():
            save_suggestion(suggestion, category, details)
            st.success(
                "¡Gracias por tu sugerencia! El equipo de IGL la revisará pronto."
            )
        else:
            st.error("Por favor, escribe una sugerencia antes de enviar.")


# File Upload Section
st.sidebar.markdown("### 📂 Subir un archivo de datos")
with st.container():
    uploaded_file = st.sidebar.file_uploader(
        "Carga tu conjunto de datos (XLSX)", type="xlsx"
    )
    if uploaded_file:
        st.session_state.uploaded_file = uploaded_file
        df = load_data(st.session_state.uploaded_file)

        data_tab1, data_tab2, data_tab3 = st.tabs(
            [
                "📋 Detalles del Reporte",
                "📊 Vista previa de datos",
                "📑 Variables por sección",
            ]
        )

        with data_tab1:
            st.markdown("### Información del Centro ZASCA")
            cohort_details = {
                "centro": st.text_input(
                    "Nombre del Centro ZASCA", placeholder="Ej.: Ciudad Bolívar"
                ),
                "cohorte": st.text_input(
                    "Número de la Cohorte", placeholder="Ej.: Primera Cohorte"
                ),
                "sector": st.text_input(
                    "Sector", placeholder="Ej.: Textiles, Manufactura, Servicios"
                ),
                "informacion_adicional": st.text_input(
                    "Información adicional", placeholder="Ej.: Fechas"
                ),
            }
            COHORT_INFO_STR = ", ".join(
                f"{key}: {value}" for key, value in cohort_details.items() if value
            )

        with data_tab2:
            st.markdown("### Vista previa de los datos cargados")
            st.dataframe(df, use_container_width=True)

        with data_tab3:
            for section_title, variables in sections_config.items():
                with st.expander(f"### {section_title}"):
                    for var_config in variables:
                        var_pair, var_type, metadata = var_config
                        initial, final = var_pair

                        # Format the variable names
                        if isinstance(initial, list):
                            initial_str = ", ".join(initial)
                        elif initial is None:
                            initial_str = "N/A"
                        else:
                            initial_str = initial

                        st.markdown(
                            f"- **{metadata['description']}**\n"
                            f"  - Inicial: `{initial_str}`\n"
                            f"  - Final: `{final}`\n"
                            f"  - Tipo: `{var_type}`"
                        )

        # Sidebar controls
        st.sidebar.markdown("---")
        model_name = st.sidebar.selectbox(
            "Modelo de OpenAI", ["gpt-3.5-turbo", "gpt-4o-2024-08-06"], index=0
        )

        st.sidebar.markdown("<br>", unsafe_allow_html=True)
        generate_report = st.sidebar.button(
            "🚀 Generar Reporte",
            help="Haz clic para generar el reporte basado en los datos y detalles proporcionados",
            use_container_width=True,
        )

        # Report Generation Process
        if generate_report:
            # Create progress indicators in sidebar
            with st.sidebar:
                progress_bar = st.progress(0)
                status_text = st.empty()

            # Generate sections and get missing variables
            report_sections, missing_variables = aggregate_data(df, sections_config)

            # Show warning in sidebar if variables are missing
            if missing_variables:
                with st.sidebar:
                    st.markdown("### Variables no encontradas")
                    st.warning(
                        "Las siguientes variables no fueron encontradas en los datos:",
                        icon="⚠️",
                    )
                    for var in missing_variables:
                        st.markdown(f"- {var}")

            # Continue with report generation
            with st.sidebar:
                status_text.info("🤖 Generando contenido de secciones...")
            asyncio.run(
                generate_section_contents(
                    report_sections, COHORT_INFO_STR, model_name, progress_bar
                )
            )

            with st.sidebar:
                status_text.info("📝 Generando resumen ejecutivo...")
            resumen_ejecutivo = asyncio.run(
                generate_executive_summary(report_sections, COHORT_INFO_STR, model_name)
            )

            with st.sidebar:
                status_text.info("✍️ Realizando edición final...")
            edited_output = asyncio.run(
                edit_report_sections(report_sections, model_name)
            )

            with st.sidebar:
                status_text.info("💾 Preparando archivo JSON...")
            json_str = generate_json_output(report_sections, resumen_ejecutivo)

            # Show success in sidebar and clear progress
            with st.sidebar:
                status_text.empty()
                progress_bar.empty()
                st.session_state.success_message = "🎉 ¡Reporte generado exitosamente!"

            # Store in session state at the end
            st.session_state.report_generated = True
            st.session_state.report_finalized = True
            st.session_state.json_str = json_str
            st.session_state.resumen_ejecutivo = resumen_ejecutivo
            st.session_state.edited_output = edited_output
            st.session_state.report_sections = report_sections
            st.session_state.markdown_content = f"""# Reporte ZASCA\n\n## Resumen Ejecutivo\n{resumen_ejecutivo}\n\n{edited_output}"""

    # Display download buttons and results (outside generate_report block)
    if st.session_state.report_finalized:
        # Show persistent success message
        if st.session_state.success_message:
            st.sidebar.success(st.session_state.success_message)

        # Create Word document
        def create_word_doc(edited=True):
            """Create a Word document with the report content"""
            doc = Document()

            # Add title
            title = doc.add_heading("Reporte ZASCA", 0)
            title.alignment = 1  # Center alignment

            # Add executive summary
            doc.add_heading("Resumen Ejecutivo", level=1)
            doc.add_paragraph(st.session_state.resumen_ejecutivo)

            def process_paragraph_text(text):
                """Process a paragraph text to handle markdown bold"""
                parts = text.split('**')
                if len(parts) <= 1:  # No bold markers
                    return doc.add_paragraph(text)
                
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if not part:  # Skip empty strings between consecutive **
                        continue
                    # Even indices (0, 2, 4, ...) are normal text, odd indices are bold
                    run = p.add_run(part)
                    run.bold = (i % 2) == 1

            if edited:
                # Process edited content by splitting into lines and handling headers
                lines = st.session_state.edited_output.split('\n')
                current_paragraph = []
                
                for line in lines:
                    if line.startswith('###'):
                        # If we have accumulated text, add it as a paragraph
                        if current_paragraph:
                            process_paragraph_text('\n\n'.join(current_paragraph))
                            current_paragraph = []
                        
                        # Add ### header as level 3
                        doc.add_heading(line.replace('###', '').strip(), level=3)
                    elif line.startswith('##'):
                        # If we have accumulated text, add it as a paragraph
                        if current_paragraph:
                            process_paragraph_text('\n\n'.join(current_paragraph))
                            current_paragraph = []
                        
                        # Add ## header as level 2
                        doc.add_heading(line.replace('##', '').strip(), level=2)
                    else:
                        # Accumulate non-header lines
                        if line.strip():  # Only add non-empty lines
                            current_paragraph.append(line)
                
                # Add any remaining text as final paragraph
                if current_paragraph:
                    process_paragraph_text('\n\n'.join(current_paragraph))
            else:
                # Add unedited content with variables subsections
                for report_section in st.session_state.report_sections:
                    # Process section content for headers if it exists
                    if report_section.content:
                        lines = report_section.content.split('\n')
                        current_paragraph = []
                        
                        for line in lines:
                            if line.startswith('###'):
                                if current_paragraph:
                                    process_paragraph_text('\n\n'.join(current_paragraph))
                                    current_paragraph = []
                                doc.add_heading(line.replace('###', '').strip(), level=3)
                            elif line.startswith('##'):
                                if current_paragraph:
                                    process_paragraph_text('\n\n'.join(current_paragraph))
                                    current_paragraph = []
                                doc.add_heading(line.replace('##', '').strip(), level=2)
                            else:
                                if line.strip():
                                    current_paragraph.append(line)
                        
                        if current_paragraph:
                            process_paragraph_text('\n\n'.join(current_paragraph))

                    # Add variables subsection
                    if report_section.variables:
                        doc.add_heading("Variables Analizadas", level=2)
                        for _, variable_data in report_section.variables.items():
                            p = doc.add_paragraph()
                            # Add variable description and interpretation
                            p.add_run(f"{variable_data.description}: ").bold = True
                            p.add_run(variable_data.interpretation)

            # Save to bytes buffer
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            return docx_buffer

        # Download buttons
        st.sidebar.markdown("### 💾 Descargar Reporte")
        col1, col2 = st.sidebar.columns(2)
        with col1:
            unedited_docx = create_word_doc(edited=False)
            st.download_button(
                label="📊 Sin Editar - Word",
                data=unedited_docx,
                file_name="reporte_zasca_sin_editar.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Descarga el reporte sin editar en formato Word",
            )
        with col2:
            edited_docx = create_word_doc(edited=True)
            st.download_button(
                label="📝 Editado - Word",
                data=edited_docx,
                file_name="reporte_zasca_editado.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Descarga el reporte editado en formato Word",
            )

        # Results display
        result_tab1, result_tab2 = st.tabs(
            ["📝 Reporte Final Editado", "🔄 Secciones Sin Editar"]
        )

        with result_tab1:
            st.markdown("### Reporte Final")
            st.markdown("#### Resumen Ejecutivo")
            st.markdown(st.session_state.resumen_ejecutivo)
            st.markdown(st.session_state.edited_output)

        with result_tab2:
            st.markdown("### Contenido Sin Editar por Sección")
            for section in st.session_state.report_sections:
                with st.expander(f"📑 {section.title}"):
                    st.markdown("#### Contenido generado")
                    st.markdown(section.content)
                    st.markdown("#### Variables procesadas")
                    for var_name, var_data in section.variables.items():
                        st.markdown(
                            f"**{var_data.description}**: {var_data.interpretation}"
                        )
                        st.markdown("---")


# Footer with Logos
st.markdown("---")
_, col1, _, col3, _ = st.columns([1, 2, 3, 2, 1])

with col1:
    st.image("assets/innpulsa_logo.png", use_container_width=True)

with col3:
    st.image("assets/igl_logo.png", use_container_width=True)
