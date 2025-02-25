"""Utilities for document generation."""

import io
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def process_paragraph_text(doc: Document, text: str) -> None:
    """
    Process markdown text into a Word paragraph with proper formatting.

    Args:
        doc: Word document instance
        text: Text to process, may contain markdown bold markers (**)
    """
    parts = text.split("**")
    if len(parts) <= 1:  # No bold markers
        doc.add_paragraph(text)
        return

    p = doc.add_paragraph()
    for i, part in enumerate(parts):
        if not part:  # Skip empty strings between consecutive **
            continue
        run = p.add_run(part)
        run.bold = (i % 2) == 1
        run.font.size = Pt(11)


def process_content_lines(doc: Document, lines: List[str]) -> None:
    """
    Process content lines into Word document with proper header levels.

    Args:
        doc: Word document instance
        lines: List of text lines to process
    """
    current_paragraph = []

    for line in lines:
        if line.startswith("###"):
            if current_paragraph:
                process_paragraph_text(doc, "\n\n".join(current_paragraph))
                current_paragraph = []

            heading = doc.add_heading(line.replace("###", "").strip(), level=3)
            heading.style.font.size = Pt(12)
        elif line.startswith("##"):
            if current_paragraph:
                process_paragraph_text(doc, "\n\n".join(current_paragraph))
                current_paragraph = []

            heading = doc.add_heading(line.replace("##", "").strip(), level=2)
            heading.style.font.size = Pt(13)
        else:
            if line.strip():  # Only add non-empty lines
                current_paragraph.append(line)

    if current_paragraph:
        process_paragraph_text(doc, "\n\n".join(current_paragraph))


def create_word_doc(session_state: Dict[str, Any], edited: bool = True) -> io.BytesIO:
    """
    Create a Word document from the report content.

    Args:
        session_state: Streamlit session state containing report data
        edited: Whether to use edited or unedited content

    Returns:
        BytesIO buffer containing the Word document
    """
    doc = Document()

    # Add title
    title = doc.add_heading("Reporte ZASCA", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add executive summary
    doc.add_heading("Resumen Ejecutivo", level=1)
    doc.add_paragraph(session_state.resumen_ejecutivo)

    if edited:
        # Process edited content
        lines = session_state.edited_output.split("\n")
        process_content_lines(doc, lines)
    else:
        # Add unedited content with variables subsections
        for report_section in session_state.report_sections:
            doc.add_heading(report_section.title, level=2)

            if report_section.content:
                lines = report_section.content.split("\n")
                process_content_lines(doc, lines)

            if report_section.variables:
                doc.add_heading("Variables Analizadas", level=2)
                for _, variable_data in report_section.variables.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{variable_data.description}: ").bold = True
                    p.add_run(variable_data.interpretation)

    # Save to bytes buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer
